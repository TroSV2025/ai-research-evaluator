import pdfplumber
from docx import Document
import io

def read_file(file):
    """
    Đọc nội dung từ file .pdf hoặc .docx
    Hỗ trợ cả:
    - file upload (Flask FileStorage)
    - đường dẫn file (string)
    """

    try:
        # ================= PDF =================
        if hasattr(file, "filename") and file.filename.endswith(".pdf"):
            text = ""
            with pdfplumber.open(file.stream) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            return text

        elif isinstance(file, str) and file.endswith(".pdf"):
            text = ""
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            return text

        # ================= DOCX =================
        elif hasattr(file, "filename") and file.filename.endswith(".docx"):
            document = Document(file.stream)
            return "\n".join(p.text for p in document.paragraphs)

        elif isinstance(file, str) and file.endswith(".docx"):
            document = Document(file)
            return "\n".join(p.text for p in document.paragraphs)

        return ""

    except Exception as e:
        return f"Lỗi đọc file: {str(e)}"
